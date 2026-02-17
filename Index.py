import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime
import os
import json
import glob
from fpdf import FPDF

# --- PAGE CONFIG ---
st.set_page_config(
    page_title="SM MOM ",
    page_icon="Logo/Picsart_23-05-18_16-47-20-287-removebg-preview.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- DRAFT MANAGER ---
DRAFTS_DIR = "drafts"
if not os.path.exists(DRAFTS_DIR):
    os.makedirs(DRAFTS_DIR)

def save_draft(file_date, time_val, year_val, dept_val, points_data, attendance_list):
    """Saves the current session state to a JSON file."""
    # Convert numpy int64 to native python int for JSON serialization
    if isinstance(year_val, list):
        year_val = [int(y) for y in year_val]
        
    draft_data = {
        "date": str(file_date),
        "time": time_val,
        "year": year_val,
        "department": dept_val,
        "points": points_data,
        "attendance": attendance_list
    }
    filepath = os.path.join(DRAFTS_DIR, f"{file_date}.json")
    with open(filepath, "w") as f:
        json.dump(draft_data, f, indent=4)
    return filepath

def load_draft(file_date):
    """Loads a draft from a JSON file into session state."""
    filepath = os.path.join(DRAFTS_DIR, f"{file_date}.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, "r") as f:
                data = json.load(f)
            return data
        except (json.JSONDecodeError, ValueError):
            return None
    return None

def get_saved_drafts():
    """Returns a list of saved draft dates."""
    files = glob.glob(os.path.join(DRAFTS_DIR, "*.json"))
    dates = [os.path.basename(f).replace(".json", "") for f in files]
    dates.sort(reverse=True)
    return dates

# --- CUSTOM CSS ---
st.markdown("""
    <style>
    /* Import Google Font */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

    /* Global Typography Override */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
        color: #1F2937; 
    }

    /* Headings */
    h1, h2, h3, h4, h5, h6 {
        color: #111827;
        font-weight: 700;
        letter-spacing: -0.025em;
        margin-bottom: 0.5rem;
    }
    
    h1 {
        font-size: 2.2rem;
    }
    
    /* Input Fields & Select Boxes */
    div[data-baseweb="select"] > div,
    .stTextInput input, 
    .stDateInput input, 
    .stTimeInput input, 
    .stTextArea textarea {
        background-color: #FFFFFF !important;
        border: 1px solid #D1D5DB !important;
        border-radius: 8px !important;
        color: #1F2937 !important;
        box-shadow: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
    }
    
    /* Focus states */
    div[data-baseweb="select"] > div:focus-within,
    .stTextInput input:focus, 
    .stDateInput input:focus, 
    .stTimeInput input:focus, 
    .stTextArea textarea:focus {
        border-color: #4F46E5 !important;
        box-shadow: 0 0 0 3px rgba(79, 70, 229, 0.1) !important;
    }

    /* Dropdown Menu Items */
    ul[data-testid="stSelectboxVirtualDropdown"] li {
        background-color: white !important;
        color: #1F2937 !important;
    }
    
    /* Radio Buttons & Checkboxes Labels */
    .stRadio label, .stCheckbox label {
        color: #374151 !important;
        font-weight: 500;
        font-size: 0.95rem;
    }
    
    /* Fix Streamlit's checkbox specific layout */
    [data-testid="stCheckbox"] label > div:first-child {
        background-color: white;
        border-color: #D1D5DB;
    }
    
    /* Card Styles */
    [data-testid="stVerticalBlockBorderWrapper"] {
        background-color: white;
        border-radius: 12px;
        border: 1px solid #E5E7EB;
        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1), 0 1px 2px 0 rgba(0, 0, 0, 0.06);
        padding: 1.5rem !important;
        margin-bottom: 1.5rem;
    }
    
    /* Primary Button */
    .stButton > button {
        background-color: #4F46E5 !important;
        color: white !important;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 1.5rem !important;
        font-weight: 600;
        font-size: 1rem !important;
        letter-spacing: 0.025em;
        transition: all 0.2s;
        box-shadow: 0 1px 2px 0 rgba(0, 0, 0, 0.05);
    }
    .stButton > button:hover {
        background-color: #4338CA !important;
        transform: translateY(-1px);
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
    }
    
    /* Status Tags */
    .status-badge {
        background-color: #ECFDF5;
        color: #065F46;
        padding: 4px 12px;
        border-radius: 9999px;
        font-size: 0.875rem;
        font-weight: 600;
        border: 1px solid #A7F3D0;
    }

    /* Headers inside cards */
    .card-title {
        font-size: 1.125rem;
        font-weight: 600;
        color: #111827;
        margin-bottom: 1.25rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
        border-bottom: 1px solid #F3F4F6;
        padding-bottom: 0.75rem;
    }
    
    /* Hide default footer */
    footer {visibility: hidden;}
    
    /* Custom divider */
    .divider {
        height: 1px;
        background-color: #E5E7EB;
        margin: 1.5rem 0;
    }
    
    /* Success Message Styles */
    .stAlert {
        background-color: #ECFDF5;
        color: #065F46;
        border: 1px solid #A7F3D0;
    }
    
    /* Reduce default Streamlit padding - adjusted to prevent clipping */
    .block-container {
        padding-top: 3rem !important;
        padding-bottom: 1rem !important;
    }
    </style>
""", unsafe_allow_html=True)

# --- SIDEBAR: DRAFT HISTORY ---
with st.sidebar:
    st.markdown("### 📂 Draft History")
    saved_drafts = get_saved_drafts()
    
    if saved_drafts:
        selected_draft = st.selectbox("Select Date to Load", saved_drafts)
        
        if st.button("📂 Load Selected Draft"):
            data = load_draft(selected_draft)
            if data:
                # Update session state with loaded values
                try:
                    loaded_date = datetime.strptime(data["date"], "%Y-%m-%d").date()
                    st.session_state["loaded_date"] = loaded_date
                    st.session_state["loaded_time"] = data["time"]
                    st.session_state["loaded_year"] = data["year"] # Now a list
                    st.session_state["loaded_dept"] = data["department"]
                    st.session_state["loaded_attendance"] = data["attendance"]
                    st.session_state.points = data["points"]
                    st.success(f"Loaded draft: {selected_draft}")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error loading draft: {e}")
    else:
        st.info("No saved drafts found.")
    
    st.markdown("---")
    st.info("💡 Saved drafts allow you to continue working on a meeting from where you left off.")

# --- HEADER SECTION ---
with st.container():
    col1, col2 = st.columns([1.5, 8.5])
    with col1:
        # Display Logo here (SM Logo)
        logo_path = os.path.join(os.path.dirname(__file__), "Logo", "Picsart_23-05-18_16-47-20-287-removebg-preview.png")
        if os.path.exists(logo_path):
            st.image(logo_path, width=110)
    with col2:
        # Removed negative margin that was clipping content
        st.markdown("<h1>Service Motto Volunteers</h1>", unsafe_allow_html=True)
        st.caption("Minutes of Meeting ")

# --- LOAD DATA ---
@st.cache_data
def load_data():
    try:
        df = pd.read_excel("students.xlsx")
        if "DEPARTMENT" in df.columns:
            df["DEPARTMENT"] = df["DEPARTMENT"].replace("AIDS", "AI&DS")
        if "NAME" in df.columns:
            df = df.drop_duplicates(subset=["NAME"], keep="first")
        return df
    except Exception as e:
        return pd.DataFrame()

df = load_data()

if df.empty:
    st.error("Could not load 'students.xlsx'. Please check if the file exists.")
    st.stop()

# --- MAIN CONTENT ---


# SECTION 1: MEETING DETAILS
with st.container():
    st.markdown('<div class="card-title">📅 Meeting Details</div>', unsafe_allow_html=True)
    
    # Defaults or Loaded Values
    default_date = st.session_state.get("loaded_date", date.today())
    default_time = st.session_state.get("loaded_time", "04.00 PM")

    c1, c2 = st.columns(2)
    with c1:
        meeting_date = st.date_input("Date", value=default_date)
    with c2:
        meeting_time = st.text_input("Time", value=default_time)

# SECTION 2: ATTENDANCE
with st.container():
    st.markdown('<div class="card-title">👥 Attendance & Participants</div>', unsafe_allow_html=True)
    
    # Selection Controls
    col_sel_1, col_sel_2 = st.columns(2)
    
    # Defaults
    year_options = sorted(df["YEAR"].unique()) if "YEAR" in df.columns else []
    
    # Handle loaded years (which might be a list or single value from old drafts)
    loaded_year = st.session_state.get("loaded_year")
    default_years = []
    if loaded_year:
        if isinstance(loaded_year, list):
            default_years = [y for y in loaded_year if y in year_options]
        elif loaded_year in year_options:
            default_years = [loaded_year]

    dept_options = sorted(df["DEPARTMENT"].unique()) if "DEPARTMENT" in df.columns else []
    
    # Handle loaded depts (which might be a list or single value from old drafts)
    loaded_dept = st.session_state.get("loaded_dept")
    default_depts = []
    if loaded_dept:
        if isinstance(loaded_dept, list):
            default_depts = [d for d in loaded_dept if d in dept_options]
        elif loaded_dept in dept_options:
            default_depts = [loaded_dept]

    with col_sel_1:
        # Changed to multiselect as requested
        selected_years = st.multiselect("Select Year(s)", year_options, default=default_years)
    
    with col_sel_2:
        # Changed to multiselect as requested
        selected_depts = st.multiselect("Select Department(s)", dept_options, default=default_depts)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # Filter Data
    if selected_years and selected_depts:
        filtered_df = df[(df["YEAR"].isin(selected_years)) & (df["DEPARTMENT"].isin(selected_depts))]
        filtered_df = filtered_df.drop_duplicates(subset=["NAME"], keep="first")
        
        col_header, col_count = st.columns([6, 1])
        with col_header:
            year_str = ", ".join(map(str, selected_years))
            dept_str = ", ".join(selected_depts)
            st.markdown(f"**Marking Attendance for:** <span class='status-badge'>{year_str} - {dept_str}</span>", unsafe_allow_html=True)
        with col_count:
            st.markdown(f"**Total:** {len(filtered_df)}")

        st.markdown("<br>", unsafe_allow_html=True)
        
        # Determine strict attendance state based on widget keys
        loaded_attendance = st.session_state.get("loaded_attendance", [])
        
        present_students = []
        if not filtered_df.empty:
            # Group by Year and display year-wise
            for year in sorted(selected_years):
                year_df = filtered_df[filtered_df["YEAR"] == year].sort_values("NAME")
                
                if not year_df.empty:
                    # Year header with select all
                    st.markdown(f"<div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 0.5rem 1rem; border-radius: 8px; margin: 1rem 0 0.5rem 0;'><span style='color: white; font-weight: 600; font-size: 1rem;'>Year {year} ({len(year_df)} students)</span></div>", unsafe_allow_html=True)
                    
                    # Select All checkbox for this year
                    select_all_key = f"select_all_year_{year}"
                    if select_all_key not in st.session_state:
                        st.session_state[select_all_key] = False
                    
                    select_all = st.checkbox(f"✓ Select All Year {year}", key=select_all_key)
                    
                    # Grid Layout for Checkboxes
                    check_cols = st.columns(4)
                    for idx, (_, row) in enumerate(year_df.iterrows()):
                        col_idx = idx % 4
                        
                        # Check if this student is in loaded attendance list or select all is checked
                        is_checked = select_all
                        if not is_checked and loaded_attendance:
                            if row["NAME"] in loaded_attendance:
                                is_checked = True
                        
                        with check_cols[col_idx]:
                            checkbox_result = st.checkbox(row["NAME"], value=is_checked, key=f"chk_{year}_{idx}_{row['NAME']}")
                            # Add to present_students if checkbox is checked (either manually or via select all)
                            if checkbox_result or select_all:
                                present_students.append(row)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
        else:
            st.info("No students found for this selection.")
    else:
        st.warning("Please select at least one Year and one Department.")

st.markdown("<div class='spacer'></div>", unsafe_allow_html=True)

# Row 3: Discussion Points
with st.container(border=True):
    st.markdown('<div class="card-header">📝 Discussion Points</div>', unsafe_allow_html=True)
    
    if "points" not in st.session_state:
        st.session_state.points = []

    # Display added points
    if st.session_state.points:
        st.markdown("<div style='background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 0.75rem 1.5rem; border-radius: 10px; margin: 1rem 0;'><span style='color: white; font-weight: 700; font-size: 1.1rem;'>📋 Added Points</span></div>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        
        for i, p in enumerate(st.session_state.points, 1):
            # Card container for each point
            st.markdown(f"""
                <div style='background: #F9FAFB; border-left: 4px solid #667eea; padding: 1rem 1.25rem; border-radius: 8px; margin-bottom: 1rem; box-shadow: 0 1px 3px rgba(0,0,0,0.1);'>
                    <div style='color: #667eea; font-weight: 700; font-size: 1.05rem; margin-bottom: 0.5rem;'>
                        {i}. {p['topic']}
                    </div>
                    <div style='color: #4B5563; line-height: 1.7; padding-left: 1.5rem;'>
                        {p['discussion']}
                    </div>
                </div>
            """, unsafe_allow_html=True)
        
        st.markdown("<div class='divider'></div>", unsafe_allow_html=True)

    st.markdown("### Add New Point")
    with st.form("discussion_form", clear_on_submit=True):
        col_topic, col_desc = st.columns([1, 2])
        with col_topic:
            new_topic = st.text_input("Topic Title", placeholder="e.g., Event Planning")
        with col_desc:
            new_discussion = st.text_area("Discussion Details", placeholder="Enter key points discussed...", height=100)
            
        submitted = st.form_submit_button("➕ Add Point")
        
        if submitted:
            if new_topic and new_discussion:
                st.session_state.points.append({
                    "topic": new_topic,
                    "discussion": new_discussion
                })
                st.success("Point added! Enter next point below.")
                st.rerun() # Immediate refresh
            else:
                st.warning("Please fill in both topic and discussion.")


st.markdown("<div class='spacer'></div>", unsafe_allow_html=True)

# ACTION BUTTONS AT BOTTOM
st.markdown("<br>", unsafe_allow_html=True)
col_empty1, col_save_bottom, col_gen_bottom, col_empty2 = st.columns([1, 2, 2, 1])

with col_save_bottom:
    if st.button("💾 Save Draft", type="secondary", use_container_width=True, key="save_bottom"):
        # We'll handle this after all data is collected
        st.session_state.save_requested = True

with col_gen_bottom:
    if st.button("🚀 Generate Reports", type="primary", use_container_width=True, key="gen_bottom"):
        st.session_state.generate_requested = True

st.markdown("<br>", unsafe_allow_html=True)

# Handle Save Draft Request (from bottom button)
if st.session_state.get("save_requested", False):
    # Collect data
    att_names = [s["NAME"] for s in present_students]
    
    saved_path = save_draft(
        meeting_date, 
        meeting_time, 
        selected_years, 
        selected_depts, 
        st.session_state.points, 
        att_names
    )
    st.success(f"Draft saved successfully for {meeting_date}!")
    st.session_state.save_requested = False

# Handle Generate Reports Request (from top button)
if st.session_state.get("generate_requested", False):
    try:
        # --- WORD DOCUMENT GENERATION ---
        template_path = os.path.join(os.path.dirname(__file__), "template.docx")
        doc = Document(template_path)

        # --- HEADER SECTION (3 Columns: Logo | Text | Logo) ---
        header_table = doc.add_table(rows=1, cols=3)
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.2)
        header_table.columns[1].width = Inches(4.1)
        header_table.columns[2].width = Inches(1.2)

        # SWAPPED LOGIC: Logo 2 (College) on Left, Logo 1 (SM) on Right
        
        # Left Logo (Now Brand_logo / College)
        cell_left = header_table.rows[0].cells[0]
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo2_path = os.path.join(os.path.dirname(__file__), "Logo", "Brand_logo.png")
        if os.path.exists(logo2_path):
            run = p_left.add_run()
            run.add_picture(logo2_path, width=Inches(1.0)) # Strictly 1.0 Inch

        # Center Text
        cell_center = header_table.rows[0].cells[1]
        p_center = cell_center.paragraphs[0]
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Updated Title
        run_title = p_center.add_run("SERVICE MOTTO VOLUNTEERS")
        run_title.bold = True
        run_title.font.size = Pt(16)
        run_title.font.name = 'Arial'
        run_title.font.color.rgb = RGBColor(0, 0, 0)
        p_center.add_run("\n")
        run_subtitle = p_center.add_run("MINUTES OF MEET")
        run_subtitle.bold = True
        run_subtitle.font.size = Pt(14)
        run_subtitle.font.name = 'Arial'
        run_subtitle.font.color.rgb = RGBColor(0, 0, 0)

        # Right Logo (Now SM Logo)
        cell_right = header_table.rows[0].cells[2]
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo1_path = os.path.join(os.path.dirname(__file__), "Logo", "Picsart_23-05-18_16-47-20-287-removebg-preview.png")
        if os.path.exists(logo1_path):
            run = p_right.add_run()
            run.add_picture(logo1_path, width=Inches(1.0)) # Strictly 1.0 Inch
        
        # Remove header borders
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        tbl = header_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:tblBorders><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders></w:tblPr>')
            tbl.insert(0, tblPr)

        doc.add_paragraph() 

        # --- MEETING DETAILS ---
        details_table = doc.add_table(rows=1, cols=2)
        details_table.autofit = False
        details_table.columns[0].width = Inches(3.25)
        details_table.columns[1].width = Inches(3.25)
        
        c1 = details_table.rows[0].cells[0]
        p = c1.paragraphs[0]
        r = p.add_run(f"DATE: {meeting_date.strftime('%d-%m-%Y')}")
        r.bold = True
        r.font.size = Pt(12)
        
        c2 = details_table.rows[0].cells[1]
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"TIME: {meeting_time}")
        r.bold = True
        r.font.size = Pt(12)

        tbl = details_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
             tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:tblBorders><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders></w:tblPr>')
             tbl.insert(0, tblPr)
        
        doc.add_paragraph() 

        # --- ATTENDANCE ---
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ["S.NO", "NAME", "YEAR", "DEPARTMENT"]
        for i, h_text in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h_text)
            run.bold = True
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        for i, s in enumerate(present_students, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = str(s["NAME"])
            row_cells[2].text = str(s.get("YEAR", ""))
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[3].text = str(s.get("DEPARTMENT", ""))
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph() 

        # --- DISCUSSION ---
        if st.session_state.points:
            heading = doc.add_paragraph()
            run = heading.add_run("Discussed in Today's SM Room Meeting:")
            run.bold = True
            run.font.size = Pt(12)
            
            # Create main discussion table
            disc_table = doc.add_table(rows=0, cols=2)
            disc_table.style = 'Table Grid'
            disc_table.autofit = False
            disc_table.columns[0].width = Inches(1.5)
            disc_table.columns[1].width = Inches(5.0)
            
            for i, p in enumerate(st.session_state.points, start=1):
                # Point row
                point_row = disc_table.add_row()
                point_cell = point_row.cells[0]
                point_para = point_cell.paragraphs[0]
                point_run = point_para.add_run(f"Point : {i}")
                point_run.bold = True
                point_run.font.size = Pt(11)
                
                # Topic in second column
                topic_cell = point_row.cells[1]
                topic_para = topic_cell.paragraphs[0]
                topic_run = topic_para.add_run(p['topic'])
                topic_run.font.size = Pt(11)
                
                # Discussion row
                disc_row = disc_table.add_row()
                disc_label_cell = disc_row.cells[0]
                disc_label_para = disc_label_cell.paragraphs[0]
                disc_label_run = disc_label_para.add_run("Discussion")
                disc_label_run.bold = True
                disc_label_run.font.size = Pt(11)
                
                # Discussion content in second column
                disc_content_cell = disc_row.cells[1]
                disc_content_para = disc_content_cell.paragraphs[0]
                
                # Split discussion by newlines and add as bullet points
                discussion_lines = p["discussion"].split('\n')
                for line in discussion_lines:
                    if line.strip():
                        disc_content_para.add_run(f"• {line.strip()}\n")
                        disc_content_para.runs[-1].font.size = Pt(11)
            
            doc.add_paragraph()
        else:
             doc.add_paragraph("No specific discussion points recorded.")

        doc.add_paragraph("\n\n")
        
        # --- SIGNATURES ---
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.autofit = False
        sig_table.columns[0].width = Inches(3.25)
        sig_table.columns[1].width = Inches(3.25)
        p1 = sig_table.rows[0].cells[0].paragraphs[0]
        r1 = p1.add_run("CONVENER")
        r1.bold = True
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2 = sig_table.rows[0].cells[1].paragraphs[0]
        r2 = p2.add_run("PRINCIPAL")
        r2.bold = True
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        tbl = sig_table._element
        tblPr = tbl.tblPr
        if tblPr is None:
             tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:tblBorders><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders></w:tblPr>')

        docx_file_name = f"SM_MOM_{meeting_date}.docx"
        doc.save(docx_file_name)

        # --- PDF GENERATION (FPDF) ---
        pdf = FPDF()
        pdf.add_page()
        
        # Header
        
        # Swapped: Brand Logo (Left), SM Logo (Right)
        if os.path.exists(logo2_path): # College Logo Left
            pdf.image(logo2_path, 10, 10, 30) # x, y, w
        
        pdf.set_y(15)
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "SERVICE MOTTO VOLUNTEERS", 0, 1, 'C')
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, "MINUTES OF MEET", 0, 1, 'C')
        
        if os.path.exists(logo1_path): # SM Logo Right
            pdf.image(logo1_path, 170, 10, 30)
        
        pdf.ln(20)
        
        # Details
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(95, 10, f"DATE: {meeting_date.strftime('%d-%m-%Y')}", 0, 0, 'L')
        pdf.cell(95, 10, f"TIME: {meeting_time}", 0, 1, 'R')
        
        pdf.ln(5)
        
        # Attendance Table
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(20, 10, "S.NO", 1, 0, 'C', 1)
        pdf.cell(90, 10, "NAME", 1, 0, 'C', 1)
        pdf.cell(30, 10, "YEAR", 1, 0, 'C', 1)
        pdf.cell(50, 10, "DEPARTMENT", 1, 1, 'C', 1)
        
        pdf.set_font("Arial", '', 11)
        for i, s in enumerate(present_students, start=1):
            pdf.cell(20, 10, str(i), 1, 0, 'C')
            pdf.cell(90, 10, str(s["NAME"]), 1, 0, 'L')
            pdf.cell(30, 10, str(s.get("YEAR", "")), 1, 0, 'C')
            pdf.cell(50, 10, str(s.get("DEPARTMENT", "")), 1, 1, 'C')
        
        pdf.ln(10)
        
        # Discussion
        if st.session_state.points:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 10, "Discussed in Today's SM Room Meeting:", 0, 1, 'L')
            pdf.ln(5)
            
            # Table header style
            pdf.set_fill_color(240, 240, 240)
            
            for i, p in enumerate(st.session_state.points, 1):
                # Point row
                pdf.set_font("Arial", 'B', 11)
                pdf.cell(40, 10, f"Point : {i}", 1, 0, 'L', 1)
                pdf.set_font("Arial", '', 11)
                pdf.cell(150, 10, p['topic'], 1, 1, 'L')
                
                # Discussion row
                pdf.set_font("Arial", 'B', 11)
                pdf.cell(40, 10, "Discussion", 1, 0, 'L', 1)
                
                # Discussion content with bullet points
                pdf.set_font("Arial", '', 11)
                discussion_lines = p["discussion"].split('\n')
                discussion_text = '\n'.join([f"• {line.strip()}" for line in discussion_lines if line.strip()])
                
                # Get current position for multi_cell
                x_pos = pdf.get_x()
                y_pos = pdf.get_y()
                pdf.multi_cell(150, 6, discussion_text, 1, 'L')
                
                pdf.ln(2)
        
        
        pdf.ln(20)
        
        # Signatures
        pdf.set_font("Arial", 'B', 12)
        y_sig = pdf.get_y()
        pdf.cell(95, 10, "CONVENER", 0, 0, 'L')
        pdf.cell(95, 10, "PRINCIPAL", 0, 1, 'R')
        
        pdf_file_name = f"SM_MOM_{meeting_date}.pdf"
        pdf.output(pdf_file_name)

        # --- DOWNLOAD BUTTONS ---
        st.success("✅ Reports Generated Successfully!")
        st.markdown("### 📥 Download Reports")
        
        d1, d2 = st.columns(2)
        with d1:
            with open(docx_file_name, "rb") as f:
                st.download_button("📘 Download Word (DOCX)", f, docx_file_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key='d_docx', use_container_width=True)
        with d2:
            with open(pdf_file_name, "rb") as f:
                st.download_button("📕 Download PDF", f, pdf_file_name, mime="application/pdf", key='d_pdf', use_container_width=True)
        
        # Reset the flag
        st.session_state.generate_requested = False
            
    except Exception as e:
        st.error(f"Error generating report: {e}")
        st.error("Traceback details (for debugging):")
        st.exception(e)
        st.session_state.generate_requested = False
